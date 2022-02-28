from docx import Document
from docx.shared import Inches 

doc = Document()

#profile picture
doc.add_picture(
    "musharraf.png",
    width=Inches(2.0)
)

# name, phone no. and email details
name = input("What is your name : ")
phone = input("What is your phone no. ")
email = input("What is your email : ")

doc.add_paragraph(name + " | " + phone + " | " + email)

# About me
doc.add_heading("About Me")
about_me=input('Tell me about yourself ? ')
doc.add_paragraph(about_me)

# Work Experience 
doc.add_heading('Work Experience')
p = doc.add_paragraph()

company = input('Enter Company ')
from_date = input('From Date ')
to_date = input('To Date ')

p.add_run(company+' ').bold = True
p.add_run(from_date + '-' + to_date + '\n').italic = True

experience_details = input(
    'Describe your experience at ' + company + ' '
)
p.add_run(experience_details)


doc.save("my-cv.docx")
