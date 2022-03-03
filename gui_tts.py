from tkinter import *
import pyttsx3
from docx import Document

def main():
    root = Tk()
    root.geometry("750x550+150+100")

    def speakpytts(txt):
        engine = pyttsx3.init()
        engine.setProperty('rate', 120)
        engine.say(txt)
        engine.runAndWait()

    def speakLable1(e):
        speakpytts("Enter your name")

    def speakLable2(e):
        speakpytts("How old are you")

    def speakLable3(e):
        speakpytts("What do you do")
    
    def speakLable4(e):
        speakpytts("Where do you work at")

    def speakLable5(e):
        speakpytts("How was your experience")


    def speakInput():
        txt = f"Hi {nameInput.get()} you are {howInput.get()} and you are a {whatInput.get()}"
        speakpytts(txt)
        doc = Document()
        doc.add_paragraph(nameInput.get() + " | " + howInput.get() + " | " + whatInput.get())
        doc.add_heading("About Me")
        doc.add_paragraph(whatInput.get())
        doc.add_heading('Work Experience')
        p = doc.add_paragraph()
        p.add_run(compInput.get()+" ").bold = True
        p.add_run(expInput.get())

        doc.save("my-cv-GUI.docx")

    # heading
    lableHeading = Label(root, text="CV Builder", font=("MADE TOMMY", 30, "bold"))
    lableHeading.grid(columnspan=2, row=0, pady=(20,0))

    # name Input 
    lableName = Label(root, text="Enter your name", font=("MADE TOMMY", 20))
    lableName.grid(column=0, row=1, sticky=W, pady=(30,0), padx=(50, 30))

    nameInput = Entry(root, relief=FLAT, font=("MADE TOMMY", 20), width =17)
    nameInput.grid(column=1, row=1, sticky=W, pady=(30,0))
    nameInput.bind('<FocusIn>', speakLable1)

    # how input
    lableHow = Label(root, text="How old are you ?", font=("MADE TOMMY", 20))
    lableHow.grid(column=0, row=2, sticky=W, pady=(30,0), padx=(50, 30))

    howInput = Entry(root, relief=FLAT, font=("MADE TOMMY", 20), width =17)
    howInput.grid(column=1, row=2, sticky=W, pady=(30,0))
    howInput.bind('<FocusIn>', speakLable2)

    # what input
    lableWhat = Label(root, text="What do you do ?", font=("MADE TOMMY", 20))
    lableWhat.grid(column=0, row=3, sticky=W, pady=(30,0), padx=(50, 30))

    whatInput = Entry(root, relief=FLAT, font=("MADE TOMMY", 20), width =17)
    whatInput.grid(column=1, row=3, sticky=W, pady=(30,0))
    whatInput.bind('<FocusIn>', speakLable3)

    # Company input
    lableComp = Label(root, text="Where do you work at ?", font=("MADE TOMMY", 20))
    lableComp.grid(column=0, row=4, sticky=W, pady=(30,0), padx=(50, 30))

    compInput = Entry(root, relief=FLAT, font=("MADE TOMMY", 20), width =17)
    compInput.grid(column=1, row=4, sticky=W, pady=(30,0))
    compInput.bind('<FocusIn>', speakLable4)

    # exp input
    lableExp = Label(root, text="How was your experience", font=("MADE TOMMY", 20))
    lableExp.grid(column=0, row=5, sticky=W, pady=(30,0), padx=(50, 30))

    expInput = Entry(root, relief=FLAT, font=("MADE TOMMY", 20), width =17)
    expInput.grid(column=1, row=5, sticky=W, pady=(30,0))
    expInput.bind('<FocusIn>', speakLable5)

    # btn
    btnSpeak = Button(root, text="Speak", font=("MADE TOMMY", 20), command=speakInput)
    btnSpeak.grid(columnspan=2, row=6, pady=35, sticky=E)

    # def addtoDoc():
    #     doc = Document()
    #     doc.add_paragraph(inputText.get() + " | " + inputText_1.get() + " | " + inputText_2.get())
    #     doc.save("my-cv.docx")

    # addtoDoc();

    root.mainloop()

if __name__ == '__main__':
    main()